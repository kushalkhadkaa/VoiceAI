"""
Complete RAG Service — local-first, ChromaDB-backed.

Features:
  - Persistent ChromaDB vector store
  - Multi-format ingestion: PDF, DOCX, TXT, MD, HTML, CSV, JSON, URLs
  - Sentence-aware chunking with configurable overlap
  - Dual embedding backend: Ollama | sentence-transformers
  - Search modes: semantic | keyword (BM25) | hybrid (RRF fusion)
  - Cross-encoder reranking (optional, requires sentence-transformers)
  - Metadata filtering: source_type, date range, doc_id, collection
  - Query analytics logging to JSONL
  - Collection export to JSON
  - Deduplication by content hash
"""
from __future__ import annotations

import hashlib
import io
import json
import logging
import math
import re
import time
import uuid
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Domain objects
# ---------------------------------------------------------------------------

@dataclass
class KBCollection:
    id: str
    name: str
    description: str
    created_at: str
    document_count: int = 0
    chunk_count: int = 0
    tags: list[str] = field(default_factory=list)
    embedding_provider: str | None = None
    embedding_dimension: int | None = None


@dataclass
class KBDocument:
    id: str
    collection_id: str
    filename: str
    source_type: str          # "file" | "url"
    source_url: str | None
    content_hash: str
    chunk_count: int
    created_at: str
    size_bytes: int = 0
    tags: list[str] = field(default_factory=list)
    page_count: int = 0


@dataclass
class KBChunk:
    id: str
    doc_id: str
    text: str
    chunk_index: int


@dataclass
class KBSearchResult:
    chunk_id: str
    doc_id: str
    doc_name: str
    collection_id: str
    text: str
    score: float
    chunk_index: int
    source_type: str = "file"
    source_url: str | None = None
    semantic_score: float | None = None
    bm25_score: float | None = None
    rerank_score: float | None = None
    page_number: int | None = None


# ---------------------------------------------------------------------------
# RAG Service
# ---------------------------------------------------------------------------

class RAGService:
    def __init__(
        self,
        db_path: str,
        embedding_provider,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        max_results: int = 5,
        similarity_threshold: float = 0.3,
        search_mode: str = "hybrid",          # semantic | keyword | hybrid
        reranking_enabled: bool = False,
        reranking_model: str = "cross-encoder/ms-marco-MiniLM-L-6-v2",
        query_analytics: bool = True,
        chunk_strategy: str = "sentence",      # sentence | word | paragraph
    ) -> None:
        self.db_path = Path(db_path)
        self.embedding_provider = embedding_provider
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.max_results = max_results
        self.similarity_threshold = similarity_threshold
        self.search_mode = search_mode
        self.reranking_enabled = reranking_enabled
        self.reranking_model = reranking_model
        self.query_analytics = query_analytics
        self.chunk_strategy = chunk_strategy

        self._client = None
        self._meta_path = self.db_path / "collections_meta.json"
        self._analytics_path = self.db_path / "query_analytics.jsonl"
        self._bm25_index: dict[str, Any] = {}   # collection_id → {texts, ids, index}
        self._reranker = None

        self.db_path.mkdir(parents=True, exist_ok=True)
        self._meta: dict[str, dict] = self._load_meta()

    # ------------------------------------------------------------------
    # ChromaDB client
    # ------------------------------------------------------------------

    @property
    def client(self):
        if self._client is None:
            try:
                import chromadb
                self._client = chromadb.PersistentClient(path=str(self.db_path / "chroma"))
            except ImportError as exc:
                raise RuntimeError("chromadb not installed. Run: pip install chromadb") from exc
        return self._client

    def _get_chroma_collection(self, collection_id: str):
        return self.client.get_or_create_collection(
            name=collection_id,
            metadata={"hnsw:space": "cosine"},
        )

    # ------------------------------------------------------------------
    # Collection management
    # ------------------------------------------------------------------
    # SQLite Metadata & Auditing
    # ------------------------------------------------------------------

    def _audit_log(self, event: str, details: str | None = None) -> None:
        try:
            from app.database import get_db_connection
            import uuid
            conn = get_db_connection()
            timestamp = _now()
            conn.execute(
                "INSERT INTO voice_audit_log (id, timestamp, user_id, event, details) VALUES (?, ?, ?, ?, ?);",
                (str(uuid.uuid4()), timestamp, "system", event, details)
            )
            conn.commit()
            conn.close()
        except Exception as exc:
            logger.warning(f"Failed to log RAG audit: {exc}")

    def _ensure_tables(self) -> None:
        from app.database import get_db_connection
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS rag_collections (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                description TEXT,
                created_at TEXT NOT NULL,
                tags TEXT,
                embedding_provider TEXT,
                embedding_dimension INTEGER
            );
        """)
        # Backfill columns on pre-existing databases (SQLite ALTER is a no-op-safe add).
        for col_def in ("embedding_provider TEXT", "embedding_dimension INTEGER"):
            col_name = col_def.split()[0]
            try:
                cursor.execute(f"ALTER TABLE rag_collections ADD COLUMN {col_def};")
            except Exception:
                pass  # column already exists
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS rag_documents (
                id TEXT PRIMARY KEY,
                collection_id TEXT NOT NULL REFERENCES rag_collections(id) ON DELETE CASCADE,
                filename TEXT NOT NULL,
                source_type TEXT NOT NULL,
                source_url TEXT,
                content_hash TEXT NOT NULL,
                chunk_count INTEGER NOT NULL,
                created_at TEXT NOT NULL,
                size_bytes INTEGER DEFAULT 0,
                tags TEXT,
                page_count INTEGER DEFAULT 0
            );
        """)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS rag_query_analytics (
                id TEXT PRIMARY KEY,
                timestamp TEXT NOT NULL,
                query TEXT NOT NULL,
                collection_ids TEXT,
                mode TEXT,
                result_count INTEGER,
                top_score REAL,
                elapsed_ms REAL
            );
        """)
        conn.commit()
        cursor.close()
        conn.close()

    def _load_meta(self) -> dict[str, dict]:
        try:
            self._ensure_tables()
            from app.database import get_db_connection
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # Fetch all collections
            cursor.execute("SELECT * FROM rag_collections;")
            col_rows = cursor.fetchall()
            
            # Fetch all documents
            cursor.execute("SELECT * FROM rag_documents;")
            doc_rows = cursor.fetchall()
            
            cursor.close()
            conn.close()
            
            meta = {}
            for col in col_rows:
                cid = col["id"]
                try:
                    tags = json.loads(col["tags"]) if col["tags"] else []
                except Exception:
                    tags = []
                col_keys = col.keys()
                meta[cid] = {
                    "id": cid,
                    "name": col["name"],
                    "description": col["description"] or "",
                    "created_at": col["created_at"],
                    "tags": tags,
                    "embedding_provider": col["embedding_provider"] if "embedding_provider" in col_keys else None,
                    "embedding_dimension": col["embedding_dimension"] if "embedding_dimension" in col_keys else None,
                    "documents": {}
                }
                
            for doc in doc_rows:
                cid = doc["collection_id"]
                if cid not in meta:
                    continue
                did = doc["id"]
                try:
                    tags = json.loads(doc["tags"]) if doc["tags"] else []
                except Exception:
                    tags = []
                meta[cid]["documents"][did] = {
                    "id": did,
                    "filename": doc["filename"],
                    "source_type": doc["source_type"],
                    "source_url": doc["source_url"],
                    "content_hash": doc["content_hash"],
                    "chunk_count": doc["chunk_count"],
                    "created_at": doc["created_at"],
                    "size_bytes": doc["size_bytes"] or 0,
                    "tags": tags,
                    "page_count": doc["page_count"] or 0
                }
                
            return meta
        except Exception as exc:
            logger.warning(f"Failed to load RAG metadata from SQLite: {exc}")
            # Fallback to local JSON if SQLite fails (migration safety)
            if self._meta_path.exists():
                try:
                    return json.loads(self._meta_path.read_text(encoding="utf-8"))
                except Exception:
                    pass
            return {}

    def _save_meta(self) -> None:
        try:
            self._ensure_tables()
            from app.database import get_db_connection
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # Use transaction to update metadata
            cursor.execute("DELETE FROM rag_documents;")
            cursor.execute("DELETE FROM rag_collections;")
            
            for cid, col in self._meta.items():
                tags_json = json.dumps(col.get("tags") or [])
                cursor.execute(
                    "INSERT INTO rag_collections (id, name, description, created_at, tags, embedding_provider, embedding_dimension) VALUES (?, ?, ?, ?, ?, ?, ?);",
                    (
                        cid,
                        col["name"],
                        col.get("description", ""),
                        col["created_at"],
                        tags_json,
                        col.get("embedding_provider"),
                        col.get("embedding_dimension"),
                    )
                )
                
                docs = col.get("documents") or {}
                for did, doc in docs.items():
                    doc_tags_json = json.dumps(doc.get("tags") or [])
                    cursor.execute(
                        """INSERT INTO rag_documents (
                            id, collection_id, filename, source_type, source_url,
                            content_hash, chunk_count, created_at, size_bytes, tags, page_count
                        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?);""",
                        (
                            did,
                            cid,
                            doc["filename"],
                            doc["source_type"],
                            doc.get("source_url"),
                            doc["content_hash"],
                            doc["chunk_count"],
                            doc["created_at"],
                            doc.get("size_bytes", 0),
                            doc_tags_json,
                            doc.get("page_count", 0)
                        )
                    )
            conn.commit()
            cursor.close()
            conn.close()
            
            # Best-effort write to JSON file for backward compatibility/debugging
            try:
                self._meta_path.write_text(json.dumps(self._meta, indent=2, default=str), encoding="utf-8")
            except Exception:
                pass
        except Exception as exc:
            logger.warning(f"Failed to save RAG metadata to SQLite: {exc}")
            # Fallback to file write
            try:
                self._meta_path.write_text(json.dumps(self._meta, indent=2, default=str), encoding="utf-8")
            except Exception:
                pass

    def _active_provider_id(self) -> str | None:
        return getattr(self.embedding_provider, "provider_id", None)

    def _active_provider_dimension(self) -> int | None:
        return getattr(self.embedding_provider, "dimension", None)

    def create_collection(self, name: str, description: str = "", tags: list[str] | None = None) -> KBCollection:
        cid = str(uuid.uuid4())
        now = _now()
        meta = {
            "id": cid, "name": name, "description": description,
            "created_at": now, "documents": {}, "tags": tags or [],
            "embedding_provider": self._active_provider_id(),
            "embedding_dimension": self._active_provider_dimension(),
        }
        self._meta[cid] = meta
        self._save_meta()
        self._get_chroma_collection(cid)
        self._audit_log("rag_collection_created", f"Created collection '{name}' ({cid})")
        return KBCollection(id=cid, name=name, description=description, created_at=now, tags=tags or [])

    def delete_collection(self, collection_id: str) -> bool:
        if collection_id not in self._meta:
            return False
        col_name = self._meta[collection_id].get("name", collection_id)
        try:
            self.client.delete_collection(collection_id)
        except Exception:
            pass
        del self._meta[collection_id]
        self._bm25_index.pop(collection_id, None)
        self._save_meta()
        self._audit_log("rag_collection_deleted", f"Deleted collection '{col_name}' ({collection_id})")
        return True

    def list_collections(self) -> list[KBCollection]:
        result = []
        for cid, m in self._meta.items():
            docs = m.get("documents", {})
            chunk_count = sum(d.get("chunk_count", 0) for d in docs.values())
            result.append(KBCollection(
                id=cid, name=m.get("name", cid), description=m.get("description", ""),
                created_at=m.get("created_at", ""), document_count=len(docs),
                chunk_count=chunk_count, tags=m.get("tags", []),
                embedding_provider=m.get("embedding_provider"),
                embedding_dimension=m.get("embedding_dimension"),
            ))
        return result

    def get_collection(self, collection_id: str) -> KBCollection | None:
        m = self._meta.get(collection_id)
        if not m:
            return None
        docs = m.get("documents", {})
        chunk_count = sum(d.get("chunk_count", 0) for d in docs.values())
        return KBCollection(
            id=collection_id, name=m.get("name", collection_id),
            description=m.get("description", ""), created_at=m.get("created_at", ""),
            document_count=len(docs), chunk_count=chunk_count, tags=m.get("tags", []),
            embedding_provider=m.get("embedding_provider"),
            embedding_dimension=m.get("embedding_dimension"),
        )

    def get_collection_stats(self, collection_id: str) -> dict[str, Any]:
        """Detailed stats for a single collection."""
        m = self._meta.get(collection_id)
        if not m:
            return {"ok": False, "detail": "Collection not found"}
        docs = m.get("documents", {})
        file_docs = [d for d in docs.values() if d.get("source_type") == "file"]
        url_docs = [d for d in docs.values() if d.get("source_type") == "url"]
        total_chunks = sum(d.get("chunk_count", 0) for d in docs.values())
        total_bytes = sum(d.get("size_bytes", 0) for d in docs.values())
        dates = sorted([d.get("created_at", "") for d in docs.values()])
        return {
            "ok": True,
            "collection_id": collection_id,
            "name": m.get("name"),
            "description": m.get("description"),
            "created_at": m.get("created_at"),
            "document_count": len(docs),
            "file_count": len(file_docs),
            "url_count": len(url_docs),
            "chunk_count": total_chunks,
            "total_size_bytes": total_bytes,
            "avg_chunks_per_doc": round(total_chunks / len(docs), 1) if docs else 0,
            "first_ingested": dates[0] if dates else None,
            "last_ingested": dates[-1] if dates else None,
            "tags": m.get("tags", []),
        }

    # ------------------------------------------------------------------
    # Document management
    # ------------------------------------------------------------------

    def list_documents(self, collection_id: str) -> list[KBDocument]:
        m = self._meta.get(collection_id, {})
        result = []
        for doc_id, d in m.get("documents", {}).items():
            result.append(KBDocument(
                id=doc_id, collection_id=collection_id,
                filename=d.get("filename", ""), source_type=d.get("source_type", "file"),
                source_url=d.get("source_url"), content_hash=d.get("content_hash", ""),
                chunk_count=d.get("chunk_count", 0), created_at=d.get("created_at", ""),
                size_bytes=d.get("size_bytes", 0), tags=d.get("tags", []),
                page_count=d.get("page_count", 0),
            ))
        return result

    def delete_document(self, collection_id: str, doc_id: str) -> bool:
        m = self._meta.get(collection_id, {})
        if doc_id not in m.get("documents", {}):
            return False
        doc_filename = m.get("documents", {}).get(doc_id, {}).get("filename", doc_id)
        try:
            col = self._get_chroma_collection(collection_id)
            existing = col.get(where={"doc_id": doc_id})
            if existing and existing.get("ids"):
                col.delete(ids=existing["ids"])
        except Exception as exc:
            logger.warning("Could not delete chunks: %s", exc)
        del self._meta[collection_id]["documents"][doc_id]
        self._bm25_index.pop(collection_id, None)   # invalidate BM25 index
        self._save_meta()
        self._audit_log("rag_document_deleted", f"Deleted document '{doc_filename}' ({doc_id}) from collection '{collection_id}'")
        return True

    def get_document_chunks(self, collection_id: str, doc_id: str, limit: int = 10) -> list[dict]:
        try:
            col = self._get_chroma_collection(collection_id)
            result = col.get(where={"doc_id": doc_id}, limit=limit, include=["documents", "metadatas"])
            chunks = []
            for text, meta in zip(result.get("documents", []), result.get("metadatas", [])):
                chunks.append({
                    "text": text,
                    "chunk_index": meta.get("chunk_index", 0),
                    "page_number": meta.get("page_number"),
                })
            chunks.sort(key=lambda c: c["chunk_index"])
            return chunks
        except Exception as exc:
            logger.warning("get_document_chunks error: %s", exc)
            return []

    # ------------------------------------------------------------------
    # Ingestion
    # ------------------------------------------------------------------

    def ingest_file(self, collection_id: str, data: bytes, filename: str, mime_type: str | None = None, tags: list[str] | None = None) -> KBDocument:
        if collection_id not in self._meta:
            raise ValueError(f"Collection {collection_id!r} does not exist")

        content_hash = hashlib.sha256(data).hexdigest()

        # Deduplication: check if same content_hash already in collection
        existing_docs = self._meta[collection_id].get("documents", {})
        for doc_id, d in existing_docs.items():
            if d.get("content_hash") == content_hash:
                logger.info("Skipping duplicate document %s (hash match)", filename)
                return KBDocument(
                    id=doc_id, collection_id=collection_id,
                    filename=d["filename"], source_type=d["source_type"],
                    source_url=d.get("source_url"), content_hash=content_hash,
                    chunk_count=d["chunk_count"], created_at=d["created_at"],
                    size_bytes=d["size_bytes"],
                )

        text, page_count = self._extract_text(data, filename, mime_type)
        if not text.strip():
            raise ValueError(f"No text could be extracted from {filename!r}")

        doc = self._ingest_text(
            collection_id=collection_id, text=text, filename=filename,
            source_type="file", source_url=None, content_hash=content_hash,
            size_bytes=len(data), tags=tags or [], page_count=page_count,
        )
        # Keep the original file so the UI can open the exact source (PDF/docx/etc.).
        try:
            self._save_raw_file(collection_id, doc.id, filename, data)
        except Exception as exc:
            logger.warning("Could not store original KB file for %s: %s", filename, exc)
        return doc

    def _kb_files_dir(self) -> Path:
        return self.db_path.parent / "kb_files"

    def _save_raw_file(self, collection_id: str, doc_id: str, filename: str, data: bytes) -> None:
        ext = Path(filename).suffix or ".bin"
        dest_dir = self._kb_files_dir() / collection_id
        dest_dir.mkdir(parents=True, exist_ok=True)
        (dest_dir / f"{doc_id}{ext}").write_bytes(data)

    def get_raw_file_path(self, collection_id: str, doc_id: str) -> Path | None:
        d = self._kb_files_dir() / collection_id
        if not d.exists():
            return None
        for p in d.glob(f"{doc_id}.*"):
            if p.is_file():
                return p
        return None

    def ingest_url(self, collection_id: str, url: str, tags: list[str] | None = None) -> KBDocument:
        if collection_id not in self._meta:
            raise ValueError(f"Collection {collection_id!r} does not exist")

        raw, ct = _fetch_url(url, timeout=20)

        if "html" in ct or url.lower().endswith((".html", ".htm", "/")):
            text = _html_to_text(raw.decode("utf-8", errors="replace"))
        else:
            text = raw.decode("utf-8", errors="replace")

        if not text.strip():
            raise ValueError(f"No text extracted from URL {url!r}")

        content_hash = hashlib.sha256(raw).hexdigest()
        filename = _url_to_filename(url)

        return self._ingest_text(
            collection_id=collection_id, text=text, filename=filename,
            source_type="url", source_url=url, content_hash=content_hash,
            size_bytes=len(raw), tags=tags or [], page_count=0,
        )

    def crawl_site(
        self,
        collection_id: str,
        start_url: str,
        max_pages: int = 500,
        same_domain_only: bool = True,
        delay_ms: int = 150,
        tags: list[str] | None = None,
        render_js: bool = True,
    ):
        """Generator that crawls a website and ingests each page.

        Renders each page in a headless Chromium (Playwright) so JavaScript / React
        SPAs are fully painted before extraction; seeds coverage from sitemap.xml;
        extracts clean main content (header/footer/nav stripped) via trafilatura; and
        skips pages whose content duplicates an already-ingested page.

        Yields dicts with crawl progress (status/url/page/ingested/failed/done).
        """
        import time as _time
        from urllib.parse import urljoin, urlparse, urldefrag

        if collection_id not in self._meta:
            raise ValueError(f"Collection {collection_id!r} does not exist")

        parsed_start = urlparse(start_url)
        base_domain = parsed_start.netloc
        if not base_domain:
            raise ValueError(f"Invalid start URL: {start_url!r}")

        _SKIP_EXT = (
            ".jpg", ".jpeg", ".png", ".gif", ".svg", ".ico", ".webp", ".pdf",
            ".zip", ".tar", ".gz", ".mp3", ".mp4", ".wav", ".css", ".js",
            ".woff", ".woff2", ".ttf", ".xml", ".json",
        )

        # 1) Seed from sitemap (reliable for large JS sites), then the start URL.
        visited: set[str] = set()
        seen_content: set[str] = set()
        queue: list[str] = []
        try:
            for u in _discover_sitemap_urls(start_url, base_domain, limit=max_pages * 2):
                queue.append(u)
        except Exception:
            pass
        if start_url not in queue:
            queue.insert(0, start_url)
        yield {"status": "seeded", "url": start_url, "queued": len(queue),
               "render_js": render_js, "ingested": 0, "failed": 0}

        # 2) Launch a single headless browser for the whole crawl.
        pw = browser = context = None
        if render_js:
            try:
                from playwright.sync_api import sync_playwright
                pw = sync_playwright().start()
                browser = pw.chromium.launch(headless=True, args=["--no-sandbox", "--disable-dev-shm-usage"])
                context = browser.new_context(user_agent=_HEADERS["User-Agent"], viewport={"width": 1366, "height": 900})
            except Exception as exc:
                logger.warning("Playwright unavailable — falling back to static fetch: %s", exc)
                pw = browser = context = None

        ingested = 0
        failed = 0
        try:
            while queue and ingested + failed < max_pages:
                url = queue.pop(0)
                url, _ = urldefrag(url)
                if url in visited:
                    continue
                visited.add(url)

                yield {"status": "crawling", "url": url, "page": ingested + failed + 1,
                       "total_queued": len(queue) + ingested + failed + 1,
                       "ingested": ingested, "failed": failed}

                try:
                    link_hrefs: list[str] = []
                    if context is not None:
                        page = context.new_page()
                        try:
                            page.goto(url, wait_until="domcontentloaded", timeout=30000)
                            try:
                                page.wait_for_load_state("networkidle", timeout=8000)
                            except Exception:
                                pass
                            page.wait_for_timeout(600)  # let React paint the body
                            html_str = page.content()
                            link_hrefs = page.eval_on_selector_all("a[href]", "els => els.map(e => e.href)") or []
                        finally:
                            page.close()
                        raw = html_str.encode("utf-8")
                    else:
                        raw, ct = _fetch_url(url, timeout=20)
                        html_str = raw.decode("utf-8", errors="replace")
                        link_hrefs = _extract_links(html_str, url)

                    # Discover same-domain page links.
                    for link in link_hrefs:
                        link, _ = urldefrag(link)
                        p = urlparse(link)
                        if not p.scheme.startswith("http"):
                            continue
                        if same_domain_only and p.netloc != base_domain:
                            continue
                        if any(link.lower().split("?")[0].endswith(ext) for ext in _SKIP_EXT):
                            continue
                        if link not in visited and link not in queue:
                            queue.append(link)

                    text = _extract_main_content(html_str, url)
                    if not text.strip():
                        failed += 1
                        yield {"status": "skipped", "url": url, "reason": "no text extracted",
                               "ingested": ingested, "failed": failed}
                        continue

                    # Skip pages whose extracted content duplicates another page
                    # (common for SPA routes that render the same shell).
                    content_hash = hashlib.sha256(text.strip().encode("utf-8")).hexdigest()
                    if content_hash in seen_content:
                        yield {"status": "duplicate", "url": url, "reason": "same content as another page",
                               "ingested": ingested, "failed": failed}
                        continue
                    seen_content.add(content_hash)

                    doc = self._ingest_text(
                        collection_id=collection_id, text=text, filename=_url_to_filename(url),
                        source_type="url", source_url=url, content_hash=content_hash,
                        size_bytes=len(raw), tags=tags or [], page_count=0,
                    )
                    ingested += 1
                    yield {"status": "saved", "url": url, "doc_id": doc.id, "chunks": doc.chunk_count,
                           "page": ingested + failed, "total_queued": len(queue) + ingested + failed,
                           "ingested": ingested, "failed": failed}

                except Exception as exc:
                    failed += 1
                    yield {"status": "error", "url": url, "error": str(exc),
                           "ingested": ingested, "failed": failed}

                if delay_ms > 0:
                    _time.sleep(delay_ms / 1000.0)
        finally:
            try:
                if browser: browser.close()
                if pw: pw.stop()
            except Exception:
                pass

        yield {"status": "done", "start_url": start_url, "ingested": ingested,
               "failed": failed, "total_visited": len(visited),
               "rendered": context is not None}

    def _ingest_text(
        self,
        collection_id: str,
        text: str,
        filename: str,
        source_type: str,
        source_url: str | None,
        content_hash: str,
        size_bytes: int,
        tags: list[str] | None = None,
        page_count: int = 0,
    ) -> KBDocument:
        doc_id = str(uuid.uuid4())

        if self.chunk_strategy == "sentence":
            chunks = _chunk_text_sentences(text, self.chunk_size, self.chunk_overlap)
        elif self.chunk_strategy == "paragraph":
            chunks = _chunk_text_paragraphs(text, self.chunk_size)
        else:
            chunks = _chunk_text_words(text, self.chunk_size, self.chunk_overlap)

        now = _now()
        col = self._get_chroma_collection(collection_id)
        ids = [f"{doc_id}__{i}" for i in range(len(chunks))]
        metadatas = [
            {
                "doc_id": doc_id,
                "filename": filename,
                "chunk_index": i,
                "source_type": source_type,
                "source_url": source_url or "",
                "collection_id": collection_id,
                "created_at": now,
            }
            for i in range(len(chunks))
        ]
        # Auto-reindex first if the active provider's dimension diverges from what
        # this collection was previously embedded with (provider switch safety).
        self._ensure_compatible_dimension(collection_id)

        embeddings = self.embedding_provider.embed(chunks)
        col.add(ids=ids, documents=chunks, metadatas=metadatas, embeddings=embeddings)

        # Record the embedding provider + concrete dimension used for this collection.
        actual_dim = len(embeddings[0]) if embeddings and embeddings[0] else self._active_provider_dimension()
        self._meta[collection_id]["embedding_provider"] = self._active_provider_id()
        self._meta[collection_id]["embedding_dimension"] = actual_dim

        doc_meta = {
            "id": doc_id, "filename": filename, "source_type": source_type,
            "source_url": source_url, "content_hash": content_hash,
            "chunk_count": len(chunks), "created_at": now, "size_bytes": size_bytes,
            "tags": tags or [], "page_count": page_count,
        }
        if "documents" not in self._meta[collection_id]:
            self._meta[collection_id]["documents"] = {}
        self._meta[collection_id]["documents"][doc_id] = doc_meta
        self._save_meta()
        self._audit_log("rag_document_ingested", f"Ingested document '{filename}' ({doc_id}) into collection '{collection_id}' with {len(chunks)} chunks")
        self._bm25_index.pop(collection_id, None)   # invalidate BM25 cache

        return KBDocument(
            id=doc_id, collection_id=collection_id, filename=filename,
            source_type=source_type, source_url=source_url,
            content_hash=content_hash, chunk_count=len(chunks),
            created_at=now, size_bytes=size_bytes, tags=tags or [], page_count=page_count,
        )

    # ------------------------------------------------------------------
    # Re-indexing (provider/dimension migration)
    # ------------------------------------------------------------------

    def reindex_collection(self, collection_id: str) -> dict[str, Any]:
        """Re-embed every stored chunk of a collection with the CURRENT provider.

        Reads each chunk's text from Chroma, deletes + recreates the Chroma
        collection, then re-adds with fresh embeddings. Document/chunk metadata
        in SQLite is preserved. Returns counts and the new provider/dimension.
        """
        if collection_id not in self._meta:
            return {"ok": False, "collection_id": collection_id, "detail": "Collection not found"}

        provider_id = self._active_provider_id()
        try:
            col = self._get_chroma_collection(collection_id)
            stored = col.get(include=["documents", "metadatas"])
        except Exception as exc:
            return {"ok": False, "collection_id": collection_id, "detail": f"Could not read chunks: {exc}"}

        ids = stored.get("ids", []) or []
        texts = stored.get("documents", []) or []
        metas = stored.get("metadatas", []) or []

        # Drop and recreate the Chroma collection so the new (possibly different
        # dimension) embeddings can be added cleanly.
        try:
            self.client.delete_collection(collection_id)
        except Exception:
            pass
        col = self._get_chroma_collection(collection_id)

        new_dim: int | None = self._active_provider_dimension()
        reembedded = 0
        if texts:
            try:
                # Batch to keep request sizes sane on cloud providers.
                batch = 64
                for start in range(0, len(texts), batch):
                    chunk_texts = texts[start:start + batch]
                    chunk_ids = ids[start:start + batch]
                    chunk_metas = metas[start:start + batch]
                    embeddings = self.embedding_provider.embed(chunk_texts)
                    if embeddings and embeddings[0]:
                        new_dim = len(embeddings[0])
                    col.add(ids=chunk_ids, documents=chunk_texts, metadatas=chunk_metas, embeddings=embeddings)
                    reembedded += len(chunk_texts)
            except Exception as exc:
                return {"ok": False, "collection_id": collection_id, "detail": f"Re-embedding failed: {exc}"}

        self._meta[collection_id]["embedding_provider"] = provider_id
        self._meta[collection_id]["embedding_dimension"] = new_dim
        self._save_meta()
        self._bm25_index.pop(collection_id, None)
        self._audit_log(
            "rag_collection_reindexed",
            f"Reindexed collection '{collection_id}' with {provider_id} (dim={new_dim}); {reembedded} chunks",
        )
        return {
            "ok": True,
            "collection_id": collection_id,
            "embedding_provider": provider_id,
            "embedding_dimension": new_dim,
            "chunks_reembedded": reembedded,
        }

    def reindex_all(self) -> dict[str, Any]:
        """Re-embed every collection with the current provider."""
        results = [self.reindex_collection(cid) for cid in list(self._meta.keys())]
        return {
            "ok": all(r.get("ok") for r in results) if results else True,
            "collection_count": len(results),
            "total_chunks_reembedded": sum(r.get("chunks_reembedded", 0) for r in results),
            "results": results,
        }

    def _probe_chroma_dimension(self, collection_id: str) -> int | None:
        """Read one stored embedding from Chroma to learn the collection's vector width."""
        try:
            col = self._get_chroma_collection(collection_id)
            sample = col.get(limit=1, include=["embeddings"])
            embs = sample.get("embeddings")
            if embs is not None and len(embs) > 0 and embs[0] is not None:
                return len(embs[0])
        except Exception as exc:
            logger.warning("Could not probe Chroma dimension for %s: %s", collection_id, exc)
        return None

    def _ensure_compatible_dimension(self, collection_id: str) -> None:
        """If the active provider dimension differs from the collection's stored
        dimension, auto-reindex that collection once before proceeding.

        Best-effort: any failure here is logged and swallowed so search/ingest can
        still attempt to continue rather than hard-breaking.
        """
        m = self._meta.get(collection_id)
        if not m:
            return
        stored_dim = m.get("embedding_dimension")
        active_dim = self._active_provider_dimension()
        if not active_dim:
            return
        # If metadata never recorded a dimension (legacy collection), probe the
        # actual embedding width from a stored Chroma vector so a provider switch
        # is still detected and migrated.
        if not stored_dim and m.get("documents"):
            stored_dim = self._probe_chroma_dimension(collection_id)
            if stored_dim:
                m["embedding_dimension"] = stored_dim
        # Only act when both dimensions are known and differ, and the collection
        # actually has chunks to migrate.
        if not stored_dim or stored_dim == active_dim:
            return
        if not m.get("documents"):
            # No chunks yet — just adopt the active provider/dimension.
            m["embedding_provider"] = self._active_provider_id()
            m["embedding_dimension"] = active_dim
            return
        logger.warning(
            "KB collection %s dimension mismatch (stored=%s, active=%s); auto-reindexing.",
            collection_id, stored_dim, active_dim,
        )
        try:
            self.reindex_collection(collection_id)
        except Exception as exc:
            logger.warning("Auto-reindex of collection %s failed: %s", collection_id, exc)

    # ------------------------------------------------------------------
    # Search — semantic
    # ------------------------------------------------------------------

    def _semantic_search(
        self,
        query_text: str,
        collection_ids: list[str],
        n: int,
        source_type_filter: str | None = None,
        doc_id_filter: str | None = None,
    ) -> list[KBSearchResult]:
        # Auto-reindex any target collection whose stored embedding dimension no
        # longer matches the active provider, so a provider switch never hard-breaks
        # search (query embedding dimension must equal the stored chunk dimension).
        for cid in collection_ids:
            self._ensure_compatible_dimension(cid)

        query_embedding = self.embedding_provider.embed_one(query_text)
        all_results: list[KBSearchResult] = []

        for cid in collection_ids:
            if cid not in self._meta:
                continue
            m = self._meta[cid]
            if not m.get("documents"):
                continue
            try:
                col = self._get_chroma_collection(cid)
                count = col.count()
                if count == 0:
                    continue

                conds = []
                if source_type_filter:
                    conds.append({"source_type": source_type_filter})
                if doc_id_filter:
                    conds.append({"doc_id": doc_id_filter})
                where_filter = conds[0] if len(conds) == 1 else ({"$and": conds} if conds else None)

                query_kwargs: dict[str, Any] = {
                    "query_embeddings": [query_embedding],
                    "n_results": min(n * 3, count),
                }
                if where_filter:
                    query_kwargs["where"] = where_filter

                res = col.query(**query_kwargs)
                ids = res.get("ids", [[]])[0]
                docs_texts = res.get("documents", [[]])[0]
                metadatas_list = res.get("metadatas", [[]])[0]
                distances = res.get("distances", [[]])[0]

                for chunk_id, chunk_text, meta, dist in zip(ids, docs_texts, metadatas_list, distances):
                    score = max(0.0, 1.0 - (dist / 2.0)) if dist is not None else 0.0
                    if score < self.similarity_threshold:
                        continue
                    doc_id = meta.get("doc_id", "")
                    doc_meta = m.get("documents", {}).get(doc_id, {})
                    all_results.append(KBSearchResult(
                        chunk_id=chunk_id, doc_id=doc_id,
                        doc_name=doc_meta.get("filename", ""),
                        collection_id=cid, text=chunk_text,
                        score=round(score, 4), chunk_index=meta.get("chunk_index", 0),
                        source_type=meta.get("source_type", "file"),
                        source_url=meta.get("source_url") or None,
                        semantic_score=round(score, 4),
                    ))
            except Exception as exc:
                logger.warning("Semantic search failed on collection %s: %s", cid, exc)

        all_results.sort(key=lambda r: r.score, reverse=True)
        return all_results[:n]

    # ------------------------------------------------------------------
    # Search — BM25 keyword
    # ------------------------------------------------------------------

    def _build_bm25_index(self, collection_id: str) -> dict[str, Any] | None:
        """Build or retrieve BM25 index for a collection."""
        if collection_id in self._bm25_index:
            return self._bm25_index[collection_id]

        try:
            from rank_bm25 import BM25Okapi
        except ImportError:
            return None  # gracefully degrade if not installed

        m = self._meta.get(collection_id, {})
        if not m.get("documents"):
            return None

        try:
            col = self._get_chroma_collection(collection_id)
            result = col.get(include=["documents", "metadatas"])
            ids = result.get("ids", [])
            texts = result.get("documents", [])
            metas = result.get("metadatas", [])
            if not texts:
                return None

            tokenized = [_tokenize(t) for t in texts]
            index = BM25Okapi(tokenized)
            self._bm25_index[collection_id] = {
                "index": index, "ids": ids, "texts": texts,
                "metas": metas, "tokenized": tokenized,
            }
            return self._bm25_index[collection_id]
        except Exception as exc:
            logger.warning("BM25 index build failed: %s", exc)
            return None

    def _keyword_search(
        self,
        query_text: str,
        collection_ids: list[str],
        n: int,
        source_type_filter: str | None = None,
    ) -> list[KBSearchResult]:
        all_results: list[KBSearchResult] = []

        for cid in collection_ids:
            idx_data = self._build_bm25_index(cid)
            if not idx_data:
                continue

            m = self._meta.get(cid, {})
            query_tokens = _tokenize(query_text)
            scores = idx_data["index"].get_scores(query_tokens)
            max_score = max(scores) if len(scores) > 0 else 1.0
            if max_score == 0:
                continue

            ranked = sorted(enumerate(scores), key=lambda x: x[1], reverse=True)
            for idx_pos, raw_score in ranked[:n * 3]:
                if raw_score <= 0:
                    continue
                normalized = raw_score / max_score
                chunk_id = idx_data["ids"][idx_pos]
                chunk_text = idx_data["texts"][idx_pos]
                meta = idx_data["metas"][idx_pos]

                if source_type_filter and meta.get("source_type") != source_type_filter:
                    continue

                doc_id = meta.get("doc_id", "")
                doc_meta = m.get("documents", {}).get(doc_id, {})
                all_results.append(KBSearchResult(
                    chunk_id=chunk_id, doc_id=doc_id,
                    doc_name=doc_meta.get("filename", ""),
                    collection_id=cid, text=chunk_text,
                    score=round(normalized, 4), chunk_index=meta.get("chunk_index", 0),
                    source_type=meta.get("source_type", "file"),
                    source_url=meta.get("source_url") or None,
                    bm25_score=round(normalized, 4),
                ))

        all_results.sort(key=lambda r: r.score, reverse=True)
        return all_results[:n]

    # ------------------------------------------------------------------
    # Search — hybrid (RRF fusion)
    # ------------------------------------------------------------------

    def _hybrid_search_rrf(
        self,
        query_text: str,
        collection_ids: list[str],
        n: int,
        source_type_filter: str | None = None,
        rrf_k: int = 60,
    ) -> list[KBSearchResult]:
        semantic = self._semantic_search(query_text, collection_ids, n * 2, source_type_filter)
        keyword = self._keyword_search(query_text, collection_ids, n * 2, source_type_filter)

        if not keyword:
            return semantic[:n]

        # Build rank maps by chunk_id
        sem_rank: dict[str, int] = {r.chunk_id: i + 1 for i, r in enumerate(semantic)}
        kw_rank: dict[str, int] = {r.chunk_id: i + 1 for i, r in enumerate(keyword)}

        all_chunk_ids = set(sem_rank.keys()) | set(kw_rank.keys())

        # RRF score: sum of 1/(k + rank) for each list
        rrf_scores: dict[str, float] = {}
        for cid in all_chunk_ids:
            score = 0.0
            if cid in sem_rank:
                score += 1.0 / (rrf_k + sem_rank[cid])
            if cid in kw_rank:
                score += 1.0 / (rrf_k + kw_rank[cid])
            rrf_scores[cid] = score

        # Merge result objects (prefer semantic metadata since it has embeddings)
        chunk_map: dict[str, KBSearchResult] = {r.chunk_id: r for r in semantic}
        for r in keyword:
            if r.chunk_id not in chunk_map:
                chunk_map[r.chunk_id] = r
            else:
                chunk_map[r.chunk_id].bm25_score = r.bm25_score

        sorted_ids = sorted(all_chunk_ids, key=lambda cid: rrf_scores[cid], reverse=True)
        results = []
        for cid in sorted_ids[:n]:
            r = chunk_map[cid]
            r.score = round(rrf_scores[cid], 6)
            results.append(r)

        return results

    # ------------------------------------------------------------------
    # Reranking
    # ------------------------------------------------------------------

    def _get_reranker(self):
        if self._reranker is not None:
            return self._reranker
        try:
            from sentence_transformers import CrossEncoder
            self._reranker = CrossEncoder(self.reranking_model)
            logger.info("Cross-encoder reranker loaded: %s", self.reranking_model)
        except ImportError:
            logger.warning("sentence-transformers not installed; reranking disabled")
            self._reranker = False
        except Exception as exc:
            logger.warning("Could not load reranker %s: %s", self.reranking_model, exc)
            self._reranker = False
        return self._reranker

    def _rerank(self, query: str, results: list[KBSearchResult]) -> list[KBSearchResult]:
        reranker = self._get_reranker()
        if not reranker:
            return results
        try:
            pairs = [(query, r.text) for r in results]
            scores = reranker.predict(pairs)
            for r, s in zip(results, scores):
                r.rerank_score = float(s)
            results.sort(key=lambda r: r.rerank_score or 0, reverse=True)
            # Re-normalize rerank score to 0-1 range for display
            if results:
                max_s = max(r.rerank_score or 0 for r in results) or 1.0
                min_s = min(r.rerank_score or 0 for r in results)
                for r in results:
                    raw = r.rerank_score or 0
                    r.score = round((raw - min_s) / max(max_s - min_s, 1e-9), 4)
        except Exception as exc:
            logger.warning("Reranking failed: %s", exc)
        return results

    # ------------------------------------------------------------------
    # Main query entrypoint
    # ------------------------------------------------------------------

    def query(
        self,
        query_text: str,
        collection_ids: list[str] | None = None,
        n_results: int | None = None,
        mode: str | None = None,
        source_type_filter: str | None = None,
        rerank: bool | None = None,
        doc_id_filter: str | None = None,
    ) -> list[KBSearchResult]:
        n = n_results or self.max_results
        collections = collection_ids or list(self._meta.keys())
        if not collections:
            return []

        search_mode = mode or self.search_mode
        # Document-scoped retrieval needs the chroma where-filter, which only the
        # semantic path applies — force semantic mode when a single doc is targeted.
        if doc_id_filter:
            search_mode = "semantic"
        do_rerank = rerank if rerank is not None else self.reranking_enabled

        t0 = time.time()

        if search_mode == "keyword":
            results = self._keyword_search(query_text, collections, n, source_type_filter)
        elif search_mode == "hybrid":
            results = self._hybrid_search_rrf(query_text, collections, n, source_type_filter)
        else:
            results = self._semantic_search(query_text, collections, n, source_type_filter, doc_id_filter)

        if do_rerank and results:
            results = self._rerank(query_text, results)

        results = results[:n]

        if self.query_analytics:
            self._log_query(query_text, collection_ids, results, time.time() - t0, search_mode)

        return results

    def query_advanced(
        self,
        query_text: str,
        collection_ids: list[str] | None = None,
        n_results: int = 5,
        mode: str = "hybrid",
        source_type_filter: str | None = None,
        rerank: bool = False,
        min_score: float | None = None,
        doc_id_filter: str | None = None,
    ) -> dict[str, Any]:
        """Advanced query with filters, metadata, and mode control."""
        t0 = time.time()

        old_threshold = self.similarity_threshold
        if min_score is not None:
            self.similarity_threshold = min_score

        try:
            results = self.query(
                query_text, collection_ids, n_results,
                mode=mode, source_type_filter=source_type_filter, rerank=rerank,
            )
        finally:
            self.similarity_threshold = old_threshold

        # Filter by doc_id if specified
        if doc_id_filter:
            results = [r for r in results if r.doc_id == doc_id_filter]

        elapsed_ms = round((time.time() - t0) * 1000, 1)
        return {
            "ok": True,
            "query": query_text,
            "mode": mode,
            "reranked": rerank,
            "elapsed_ms": elapsed_ms,
            "result_count": len(results),
            "results": [_result_to_dict(r) for r in results],
            "context": self._build_context_from_results(results),
        }

    def build_context(self, query_text: str, collection_ids: list[str] | None = None, doc_id: str | None = None) -> str:
        results = self.query(query_text, collection_ids, doc_id_filter=doc_id)
        return self._build_context_from_results(results)

    def _build_context_from_results(self, results: list[KBSearchResult]) -> str:
        if not results:
            return ""
        parts = []
        for r in results:
            src = r.source_url if r.source_type == "url" and r.source_url else r.doc_name
            parts.append(f"[Source: {src}]\n{r.text}")
        return "\n\n---\n\n".join(parts)

    # ------------------------------------------------------------------
    # Analytics
    # ------------------------------------------------------------------

    def _log_query(
        self,
        query: str,
        collection_ids: list[str] | None,
        results: list[KBSearchResult],
        elapsed: float,
        mode: str,
    ) -> None:
        try:
            self._ensure_tables()
            from app.database import get_db_connection
            import uuid
            conn = get_db_connection()
            ts = _now()
            col_ids_json = json.dumps(collection_ids or [])
            top_score = round(results[0].score, 4) if results else 0.0
            elapsed_ms = round(elapsed * 1000, 1)
            
            conn.execute(
                """INSERT INTO rag_query_analytics (
                    id, timestamp, query, collection_ids, mode, result_count, top_score, elapsed_ms
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?);""",
                (
                    str(uuid.uuid4()),
                    ts,
                    query,
                    col_ids_json,
                    mode,
                    len(results),
                    top_score,
                    elapsed_ms
                )
            )
            conn.commit()
            conn.close()
        except Exception as exc:
            logger.warning(f"Failed to log query analytics to SQLite: {exc}")
            
        # Best effort write to JSONL file
        try:
            entry = {
                "ts": _now(),
                "query": query,
                "collection_ids": collection_ids,
                "mode": mode,
                "result_count": len(results),
                "top_score": round(results[0].score, 4) if results else 0,
                "elapsed_ms": round(elapsed * 1000, 1),
            }
            with self._analytics_path.open("a", encoding="utf-8") as f:
                f.write(json.dumps(entry) + "\n")
        except Exception:
            pass

    def get_analytics(self, limit: int = 100) -> dict[str, Any]:
        logs: list[dict] = []
        try:
            self._ensure_tables()
            from app.database import get_db_connection
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute(
                "SELECT * FROM rag_query_analytics ORDER BY timestamp DESC LIMIT ?;",
                (limit,)
            )
            rows = cursor.fetchall()
            cursor.close()
            conn.close()
            
            for row in rows:
                try:
                    cids = json.loads(row["collection_ids"]) if row["collection_ids"] else []
                except Exception:
                    cids = []
                logs.append({
                    "ts": row["timestamp"],
                    "query": row["query"],
                    "collection_ids": cids,
                    "mode": row["mode"],
                    "result_count": row["result_count"],
                    "top_score": row["top_score"],
                    "elapsed_ms": row["elapsed_ms"]
                })
            # Reverse to match chronology expected by summary calculations
            logs.reverse()
        except Exception as exc:
            logger.warning(f"Failed to fetch query analytics from SQLite: {exc}")
            # Fallback to JSONL file reading
            logs = []
            if self._analytics_path.exists():
                try:
                    lines = self._analytics_path.read_text(encoding="utf-8").strip().splitlines()
                    for line in lines[-limit:]:
                        try:
                            logs.append(json.loads(line))
                        except Exception:
                            pass
                except Exception:
                    pass

        total = len(logs)
        avg_results = round(sum(l.get("result_count", 0) for l in logs) / total, 2) if total else 0
        avg_ms = round(sum(l.get("elapsed_ms", 0) for l in logs) / total, 1) if total else 0
        zero_results = sum(1 for l in logs if l.get("result_count", 0) == 0)

        # Top queries by frequency
        from collections import Counter
        query_counter = Counter(l.get("query", "") for l in logs)
        top_queries = [{"query": q, "count": c} for q, c in query_counter.most_common(10)]

        return {
            "ok": True,
            "total_queries": total,
            "avg_results": avg_results,
            "avg_latency_ms": avg_ms,
            "zero_result_queries": zero_results,
            "top_queries": top_queries,
            "recent": list(reversed(logs[-20:])),
        }

    def clear_analytics(self) -> None:
        try:
            self._ensure_tables()
            from app.database import get_db_connection
            conn = get_db_connection()
            conn.execute("DELETE FROM rag_query_analytics;")
            conn.commit()
            conn.close()
        except Exception as exc:
            logger.warning(f"Failed to clear query analytics from SQLite: {exc}")
            
        if self._analytics_path.exists():
            try:
                self._analytics_path.write_text("", encoding="utf-8")
            except Exception:
                pass

    # ------------------------------------------------------------------
    # Export
    # ------------------------------------------------------------------

    def export_collection(self, collection_id: str) -> dict[str, Any]:
        m = self._meta.get(collection_id)
        if not m:
            return {"ok": False, "detail": "Collection not found"}
        try:
            col = self._get_chroma_collection(collection_id)
            result = col.get(include=["documents", "metadatas"])
            chunks = []
            for text, meta in zip(result.get("documents", []), result.get("metadatas", [])):
                chunks.append({"text": text, **meta})
            return {
                "ok": True,
                "collection": {"id": collection_id, **{k: v for k, v in m.items() if k != "documents"}},
                "documents": list(m.get("documents", {}).values()),
                "chunks": chunks,
                "exported_at": _now(),
            }
        except Exception as exc:
            return {"ok": False, "detail": str(exc)}

    # ------------------------------------------------------------------
    # Status
    # ------------------------------------------------------------------

    def status(self) -> dict[str, Any]:
        try:
            emb_status = self.embedding_provider.status()
        except Exception as exc:
            emb_status = {"ok": False, "detail": str(exc)}

        collections = self.list_collections()
        total_docs = sum(c.document_count for c in collections)
        total_chunks = sum(c.chunk_count for c in collections)

        # BM25 availability
        try:
            import rank_bm25  # noqa: F401
            bm25_available = True
        except ImportError:
            bm25_available = False

        # Cross-encoder availability. Catch any error, not just ImportError:
        # sentence-transformers can be installed yet fail to import (e.g. a broken
        # torchcodec native lib raises OSError), and an optional reranker probe must
        # never take down the whole KB status endpoint.
        try:
            from sentence_transformers import CrossEncoder  # noqa: F401
            reranker_available = True
        except Exception:
            reranker_available = False

        return {
            "ok": emb_status.get("ok", False),
            "mode": "local",
            "search_mode": self.search_mode,
            "chunk_strategy": self.chunk_strategy,
            "reranking_enabled": self.reranking_enabled,
            "reranking_model": self.reranking_model,
            "db_path": str(self.db_path),
            "embedding": emb_status,
            "embedding_provider": self._active_provider_id(),
            "embedding_model": getattr(self.embedding_provider, "model_id", None),
            "embedding_dimension": self._active_provider_dimension(),
            "collection_count": len(collections),
            "document_count": total_docs,
            "chunk_count": total_chunks,
            "capabilities": {
                "bm25_available": bm25_available,
                "reranker_available": reranker_available,
                "hybrid_search": bm25_available,
            },
        }

    # ------------------------------------------------------------------
    # Text extraction
    # ------------------------------------------------------------------

    def _extract_text(self, data: bytes, filename: str, mime_type: str | None) -> tuple[str, int]:
        """Returns (text, page_count)."""
        name_lower = filename.lower()

        if name_lower.endswith(".pdf") or (mime_type and "pdf" in mime_type):
            return _extract_pdf(data)

        if name_lower.endswith(".docx") or (mime_type and "docx" in mime_type):
            return _extract_docx(data), 0

        if name_lower.endswith((".html", ".htm")) or (mime_type and "html" in mime_type):
            return _html_to_text(data.decode("utf-8", errors="replace")), 0

        if name_lower.endswith(".csv"):
            return _extract_csv(data), 0

        if name_lower.endswith(".json"):
            return _extract_json(data), 0

        return data.decode("utf-8", errors="replace"), 0


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _now() -> str:
    import datetime
    return datetime.datetime.utcnow().isoformat() + "Z"


def _tokenize(text: str) -> list[str]:
    """Simple tokenizer for BM25: lowercase + remove punctuation."""
    return re.findall(r'\b[a-z]{2,}\b', text.lower())


def _chunk_text_sentences(text: str, target_size: int, overlap_words: int) -> list[str]:
    """Sentence-aware chunking: accumulate sentences until target word count, overlap at sentence boundary."""
    # Split into sentences
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    sentences = [s.strip() for s in sentences if s.strip()]

    if not sentences:
        return []

    chunks = []
    current_sentences: list[str] = []
    current_words = 0

    for sent in sentences:
        sent_words = len(sent.split())
        if current_words + sent_words > target_size and current_sentences:
            chunks.append(" ".join(current_sentences))
            # Overlap: keep last N words worth of sentences
            overlap_sents: list[str] = []
            overlap_count = 0
            for s in reversed(current_sentences):
                w = len(s.split())
                if overlap_count + w <= overlap_words:
                    overlap_sents.insert(0, s)
                    overlap_count += w
                else:
                    break
            current_sentences = overlap_sents
            current_words = overlap_count

        current_sentences.append(sent)
        current_words += sent_words

    if current_sentences:
        chunks.append(" ".join(current_sentences))

    return [c for c in chunks if c.strip()]


def _chunk_text_paragraphs(text: str, max_words: int) -> list[str]:
    """Split by paragraphs, merge short ones, split long ones."""
    paragraphs = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip()]
    chunks = []
    current = ""
    current_words = 0

    for para in paragraphs:
        para_words = len(para.split())
        if para_words > max_words:
            if current:
                chunks.append(current.strip())
                current = ""
                current_words = 0
            # Split large paragraph into word chunks
            sub = _chunk_text_words(para, max_words, 0)
            chunks.extend(sub)
        elif current_words + para_words > max_words and current:
            chunks.append(current.strip())
            current = para
            current_words = para_words
        else:
            current = (current + "\n\n" + para).strip() if current else para
            current_words += para_words

    if current:
        chunks.append(current.strip())

    return [c for c in chunks if c.strip()]


def _chunk_text_words(text: str, chunk_size: int, overlap: int) -> list[str]:
    """Simple word-based overlapping chunks (original method)."""
    words = text.split()
    if not words:
        return []
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk = " ".join(words[start:end])
        if chunk.strip():
            chunks.append(chunk)
        if end == len(words):
            break
        start += chunk_size - overlap
    return chunks


def _extract_pdf(data: bytes) -> tuple[str, int]:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
        return "\n\n".join(parts), len(reader.pages)
    except ImportError:
        raise RuntimeError("pypdf not installed. Run: pip install pypdf")
    except Exception as exc:
        raise ValueError(f"Could not read PDF: {exc}") from exc


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        return "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    except Exception as exc:
        raise ValueError(f"Could not read DOCX: {exc}") from exc


def _extract_csv(data: bytes) -> str:
    try:
        import csv
        text_io = io.StringIO(data.decode("utf-8", errors="replace"))
        reader = csv.reader(text_io)
        rows = list(reader)
        if not rows:
            return ""
        headers = rows[0]
        lines = [", ".join(headers)]
        for row in rows[1:]:
            lines.append(", ".join(f"{h}: {v}" for h, v in zip(headers, row)))
        return "\n".join(lines)
    except Exception as exc:
        return data.decode("utf-8", errors="replace")


def _extract_json(data: bytes) -> str:
    try:
        obj = json.loads(data.decode("utf-8", errors="replace"))
        return json.dumps(obj, indent=2, ensure_ascii=False)
    except Exception:
        return data.decode("utf-8", errors="replace")


def _html_to_text(html: str) -> str:
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        # Strip scripts/styles + common boilerplate containers (header/footer/nav).
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form", "noscript", "svg"]):
            tag.decompose()
        for sel in ("[role=navigation]", "[role=banner]", "[role=contentinfo]", ".header", ".footer", ".navbar", ".nav", ".menu", ".cookie", ".breadcrumb"):
            for tag in soup.select(sel):
                tag.decompose()
        # Prefer the main content region when the page marks one.
        main = soup.find("main") or soup.find("article") or soup.find(attrs={"role": "main"}) or soup.body or soup
        text = main.get_text(separator="\n")
        return re.sub(r"\n{3,}", "\n\n", text).strip()
    except ImportError:
        text = re.sub(r"<[^>]+>", " ", html)
        return re.sub(r"\s{2,}", " ", text).strip()


def _extract_main_content(html: str, url: str | None = None) -> str:
    """Extract clean main content, dropping header/footer/nav boilerplate.

    Uses trafilatura (purpose-built for corpus/RAG extraction) when available,
    falling back to a structural BeautifulSoup strip."""
    try:
        import trafilatura
        txt = trafilatura.extract(
            html, include_comments=False, include_tables=True,
            favor_recall=True, no_fallback=False, url=url,
        )
        if txt and txt.strip():
            return txt.strip()
    except Exception:
        pass
    return _html_to_text(html)


def _discover_sitemap_urls(start_url: str, base_domain: str, limit: int = 5000) -> list[str]:
    """Seed crawl coverage from sitemap.xml / robots.txt — the reliable way to
    enumerate every page of a JS site without depending on link-following."""
    from urllib.parse import urljoin, urlparse
    found: list[str] = []
    seen_sitemaps: set[str] = set()
    queue = [urljoin(start_url, "/sitemap.xml"), urljoin(start_url, "/sitemap_index.xml"),
             urljoin(start_url, "/sitemap-index.xml")]
    try:
        raw, _ = _fetch_url(urljoin(start_url, "/robots.txt"), timeout=10)
        for m in re.findall(r"(?i)sitemap:\s*(\S+)", raw.decode("utf-8", "replace")):
            queue.append(m.strip())
    except Exception:
        pass
    queue = list(dict.fromkeys(queue))
    while queue and len(found) < limit:
        sm = queue.pop(0)
        if sm in seen_sitemaps:
            continue
        seen_sitemaps.add(sm)
        try:
            raw, _ = _fetch_url(sm, timeout=15)
            xml = raw.decode("utf-8", "replace")
        except Exception:
            continue
        for loc in re.findall(r"<loc>\s*([^<\s]+)\s*</loc>", xml):
            loc = loc.strip()
            if loc.lower().endswith(".xml"):
                if loc not in seen_sitemaps:
                    queue.append(loc)
            elif urlparse(loc).netloc == base_domain:
                found.append(loc)
    return list(dict.fromkeys(found))


def _result_to_dict(r: KBSearchResult) -> dict[str, Any]:
    return {
        "chunk_id": r.chunk_id, "doc_id": r.doc_id, "doc_name": r.doc_name,
        "collection_id": r.collection_id, "text": r.text, "score": r.score,
        "chunk_index": r.chunk_index, "source_type": r.source_type,
        "source_url": r.source_url, "semantic_score": r.semantic_score,
        "bm25_score": r.bm25_score, "rerank_score": r.rerank_score,
        "page_number": r.page_number,
    }


_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.5",
}


def _fetch_url(url: str, timeout: int = 20) -> tuple[bytes, str]:
    """Fetch a URL, follow redirects, return (raw_bytes, content_type)."""
    import urllib.request as ur
    req = ur.Request(url, headers=_HEADERS)
    try:
        with ur.urlopen(req, timeout=timeout) as resp:
            raw = resp.read()
            ct = resp.headers.get_content_type() or ""
        return raw, ct
    except Exception as exc:
        raise ValueError(f"Could not fetch {url!r}: {exc}") from exc


def _url_to_filename(url: str) -> str:
    from urllib.parse import urlparse
    p = urlparse(url)
    path = p.path.rstrip("/")
    name = path.split("/")[-1] if path else ""
    if not name:
        name = p.netloc
    # Clean up and truncate
    name = re.sub(r"[^\w\-.]", "_", name)[:80]
    return name or "webpage"


def _looks_html(url: str) -> bool:
    url_lower = url.lower().split("?")[0]
    return (
        url_lower.endswith(("/", ".html", ".htm", ".php", ".asp", ".aspx", ".jsp"))
        or "." not in url_lower.split("/")[-1]
    )


class _LinkExtractor:
    """Fast HTML link extractor using HTMLParser (no external deps)."""
    from html.parser import HTMLParser as _HP

    class _Parser(_HP):
        def __init__(self, base_url: str) -> None:
            super().__init__()
            self.base_url = base_url
            self.links: list[str] = []

        def handle_starttag(self, tag: str, attrs: list) -> None:  # type: ignore[override]
            if tag == "a":
                for k, v in attrs:
                    if k == "href" and v:
                        from urllib.parse import urljoin
                        abs_link = urljoin(self.base_url, v.strip())
                        self.links.append(abs_link)


def _extract_links(html: str, base_url: str) -> list[str]:
    """Extract all absolute href links from HTML."""
    try:
        # Try BeautifulSoup first for robustness
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        from urllib.parse import urljoin
        links = []
        for tag in soup.find_all("a", href=True):
            href = tag["href"].strip()
            if href and not href.startswith(("#", "mailto:", "tel:", "javascript:")):
                links.append(urljoin(base_url, href))
        return links
    except ImportError:
        pass
    # Fallback: regex
    from urllib.parse import urljoin
    hrefs = re.findall(r'href=["\']([^"\'#][^"\']*)["\']', html, re.IGNORECASE)
    return [urljoin(base_url, h.strip()) for h in hrefs if h.strip()]
